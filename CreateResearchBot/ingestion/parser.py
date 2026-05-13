import pathlib
import re

import PyPDF2
import docx
import openpyxl

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".txt"}


def parse_file(path: str | pathlib.Path) -> str:
    """Определяет тип файла по расширению и возвращает чистый текст."""
    p = pathlib.Path(path)
    suffix = p.suffix.lower()
    parsers = {
        ".pdf": _parse_pdf,
        ".docx": _parse_docx,
        ".xlsx": _parse_xlsx,
        ".txt": _parse_txt,
    }
    if suffix not in parsers:
        raise ValueError(
            f"Неподдерживаемый формат '{suffix}'. "
            f"Поддерживаются: {', '.join(sorted(parsers))}"
        )
    if not p.exists():
        raise FileNotFoundError(f"Файл не найден: {p}")
    return parsers[suffix](p)


# --- Внутренние парсеры ---

def _parse_pdf(path: pathlib.Path) -> str:
    parts = []
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        if reader.is_encrypted:
            raise ValueError(f"PDF защищён паролём: {path.name}")
        for page in reader.pages:
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                parts.append(text)
    return _normalize("\n\n".join(parts))


def _parse_docx(path: pathlib.Path) -> str:
    document = docx.Document(str(path))
    parts = []
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    # Таблицы тоже извлекаем
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return _normalize("\n".join(parts))


def _parse_xlsx(path: pathlib.Path) -> str:
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if cells:
                parts.append("\t".join(cells))
    wb.close()
    return _normalize("\n".join(parts))


def _parse_txt(path: pathlib.Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return _normalize(path.read_text(encoding=encoding))
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Не удалось определить кодировку файла: {path.name}")


def _normalize(text: str) -> str:
    """Убирает лишние пробелы и пустые строки, нормализует переносы."""
    # Схлопываем больше двух переносов подряд в два
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Убираем пробелы/табы в конце каждой строки
    text = "\n".join(line.rstrip() for line in text.splitlines())
    return text.strip()
